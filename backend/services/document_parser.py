"""
Document Parsing Service
Supports: PDF, DOCX, Images (OCR), Plain Text, URLs
Uses Google Document AI + Cloud Vision as primary, with fallbacks.
"""
import io
import re
from typing import Optional
import aiohttp

# Local parsers (always available)
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Google Cloud imports (optional — gracefully degrade)
try:
    from google.cloud import documentai_v1 as documentai
    HAS_DOC_AI = True
except ImportError:
    HAS_DOC_AI = False

try:
    from google.cloud import vision
    HAS_VISION = True
except ImportError:
    HAS_VISION = False

try:
    from google.cloud import translate_v2 as translate
    HAS_TRANSLATE = True
except ImportError:
    HAS_TRANSLATE = False

from backend.config import GCP_PROJECT_ID, GCP_LOCATION, DOCUMENT_AI_PROCESSOR_ID


# ============================================================
# Primary: Google Document AI (for PDFs)
# ============================================================
def parse_with_document_ai(file_content: bytes, mime_type: str = "application/pdf") -> dict:
    """Parse document using Google Document AI — the premium parser."""
    try:
        if not HAS_DOC_AI or not DOCUMENT_AI_PROCESSOR_ID or not GCP_PROJECT_ID:
            raise ValueError("Document AI not configured, falling back")
        
        client = documentai.DocumentProcessorServiceClient()
        resource_name = client.processor_path(
            GCP_PROJECT_ID, GCP_LOCATION, DOCUMENT_AI_PROCESSOR_ID
        )
        
        raw_document = documentai.RawDocument(
            content=file_content, mime_type=mime_type
        )
        
        request = documentai.ProcessRequest(
            name=resource_name, raw_document=raw_document
        )
        
        result = client.process_document(request=request)
        document = result.document
        
        entities = []
        for entity in document.entities:
            entities.append({
                "type": entity.type_,
                "text": entity.mention_text,
                "confidence": entity.confidence,
            })
        
        return {
            "text": document.text,
            "pages": len(document.pages),
            "entities": entities,
            "method": "google_document_ai",
        }
    except Exception as e:
        print(f"Document AI failed: {e}, falling back to PyMuPDF")
        return None


# ============================================================
# Fallback: PyMuPDF (for PDFs)
# ============================================================
def parse_pdf_fallback(file_content: bytes) -> dict:
    """Parse PDF using PyMuPDF as fallback."""
    if not HAS_PYMUPDF:
        return {"text": "", "pages": 0, "entities": [], "method": "no_pdf_parser"}
    doc = fitz.open(stream=file_content, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text() + "\n"
    page_count = len(doc)
    doc.close()
    
    return {
        "text": text.strip(),
        "pages": page_count,
        "entities": [],
        "method": "pymupdf_fallback",
    }


# ============================================================
# DOCX Parser
# ============================================================
def parse_docx(file_content: bytes) -> dict:
    """Parse DOCX files."""
    if not HAS_DOCX:
        return {"text": "", "pages": 0, "entities": [], "method": "no_docx_parser"}
    doc = Document(io.BytesIO(file_content))
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
    
    return {
        "text": text.strip(),
        "pages": 1,
        "entities": [],
        "method": "python_docx",
    }


# ============================================================
# Image OCR: Google Cloud Vision
# ============================================================
def parse_image_with_vision(file_content: bytes) -> dict:
    """Parse image using Google Cloud Vision API for OCR."""
    if not HAS_VISION:
        return {"text": "", "pages": 0, "entities": [], "method": "no_vision_api"}
    try:
        client = vision.ImageAnnotatorClient()
        image = vision.Image(content=file_content)
        
        response = client.document_text_detection(image=image)
        
        if response.error.message:
            raise Exception(response.error.message)
        
        text = response.full_text_annotation.text if response.full_text_annotation else ""
        
        return {
            "text": text.strip(),
            "pages": 1,
            "entities": [],
            "method": "google_cloud_vision",
        }
    except Exception as e:
        print(f"Cloud Vision failed: {e}")
        return {
            "text": "",
            "pages": 0,
            "entities": [],
            "method": "vision_failed",
        }


# ============================================================
# URL Scraper
# ============================================================
async def parse_url(url: str) -> dict:
    """Scrape text content from a URL (Terms of Service pages, etc)."""
    try:
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=aiohttp.ClientTimeout(total=15)) as response:
                html = await response.text()
        
        # Basic HTML to text (remove tags)
        text = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL)
        text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
        text = re.sub(r'<[^>]+>', '\n', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = text.strip()
        
        return {
            "text": text[:20000],  # Limit to 20k chars
            "pages": 1,
            "entities": [],
            "method": "url_scrape",
            "source_url": url,
        }
    except Exception as e:
        return {
            "text": "",
            "pages": 0,
            "entities": [],
            "method": "url_failed",
            "error": str(e),
        }


# ============================================================
# Plain Text
# ============================================================
def parse_text(text: str) -> dict:
    """Handle raw text input."""
    return {
        "text": text.strip(),
        "pages": 1,
        "entities": [],
        "method": "plain_text",
    }


# ============================================================
# Language Detection & Translation
# ============================================================
def detect_and_translate(text: str) -> dict:
    """Detect language and translate to English if needed."""
    if not HAS_TRANSLATE:
        return {"original_language": "en", "translated_text": text, "was_translated": False}
    try:
        client = translate.Client()
        
        # Detect language (use first 500 chars for detection)
        detection = client.detect_language(text[:500])
        detected_lang = detection["language"]
        confidence = detection.get("confidence", 0)
        
        if detected_lang != "en" and confidence > 0.5:
            # Translate to English (translate in chunks to handle large docs)
            chunk_size = 5000
            translated_chunks = []
            for i in range(0, len(text), chunk_size):
                chunk = text[i:i + chunk_size]
                result = client.translate(chunk, target_language="en")
                translated_chunks.append(result["translatedText"])
            
            return {
                "original_language": detected_lang,
                "translated_text": " ".join(translated_chunks),
                "was_translated": True,
                "confidence": confidence,
            }
        
        return {
            "original_language": "en",
            "translated_text": text,
            "was_translated": False,
            "confidence": confidence,
        }
    except Exception as e:
        print(f"Translation failed: {e}")
        return {
            "original_language": "unknown",
            "translated_text": text,
            "was_translated": False,
            "error": str(e),
        }


# ============================================================
# Master Parser — Routes to the right parser
# ============================================================
async def parse_document(
    file_content: Optional[bytes] = None,
    filename: Optional[str] = None,
    url: Optional[str] = None,
    raw_text: Optional[str] = None,
) -> dict:
    """
    Master document parser. Routes to the appropriate parser based on input type.
    Returns: { text, pages, entities, method, language_info }
    """
    result = None
    
    if url:
        result = await parse_url(url)
    elif raw_text:
        result = parse_text(raw_text)
    elif file_content and filename:
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
        
        if ext == "pdf":
            # Try Document AI first, fall back to PyMuPDF
            result = parse_with_document_ai(file_content, "application/pdf")
            if not result or not result.get("text"):
                result = parse_pdf_fallback(file_content)
        
        elif ext == "docx":
            result = parse_docx(file_content)
        
        elif ext in ("png", "jpg", "jpeg", "tiff", "bmp", "webp"):
            result = parse_image_with_vision(file_content)
        
        elif ext == "txt":
            result = parse_text(file_content.decode("utf-8", errors="ignore"))
        
        else:
            # Try as text
            try:
                text = file_content.decode("utf-8", errors="ignore")
                result = parse_text(text)
            except Exception:
                result = {"text": "", "pages": 0, "entities": [], "method": "unsupported"}
    
    if not result or not result.get("text"):
        return {
            "text": "",
            "pages": 0,
            "entities": [],
            "method": "failed",
            "language_info": {},
            "error": "Could not extract text from the document",
        }
    
    # Detect language and translate if needed
    language_info = detect_and_translate(result["text"])
    
    # Use translated text for analysis
    if language_info.get("was_translated"):
        result["original_text"] = result["text"]
        result["text"] = language_info["translated_text"]
    
    result["language_info"] = language_info
    
    return result
